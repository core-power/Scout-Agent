"""文档解析器 — 将各种格式的文件解析为美观的 Markdown.

支持的格式:
- PDF (PyPDF2)
- DOCX (python-docx)
- TXT / CSV / JSON
- Markdown (.md)
- HTML (.html)
- 代码文件 (.py, .js, .ts, .rs, .go 等)
"""

from __future__ import annotations

import re
import json
from pathlib import Path
from datetime import datetime
from typing import Optional


class DocumentParser:
    """文档解析器 — 统一接口."""

    SUPPORTED_EXTENSIONS = {
        ".md", ".txt", ".csv", ".json", ".html", ".htm",
        ".pdf", ".docx",
        ".py", ".js", ".ts", ".rs", ".go", ".java", ".c", ".cpp", ".h",
        ".sh", ".bash", ".yml", ".yaml", ".toml", ".ini", ".cfg",
        ".sql", ".graphql", ".xml",
    }

    @classmethod
    def parse(cls, file_path: Path) -> dict:
        """解析文件，返回结构化结果.
        
        Returns:
            {
                "title": str,           # 自动推断的标题
                "content": str,         # 解析后的 Markdown 内容
                "format": str,          # 原始格式
                "size": int,           # 文件大小 (bytes)
                "parsed_at": str,      # ISO 时间戳
                "meta": dict,          # 元数据
            }
        """
        ext = file_path.suffix.lower()
        size = file_path.stat().st_size
        
        if ext == ".pdf":
            content, meta = cls._parse_pdf(file_path)
        elif ext == ".docx":
            content, meta = cls._parse_docx(file_path)
        elif ext == ".md":
            content, meta = cls._parse_markdown(file_path)
        elif ext in (".html", ".htm"):
            content, meta = cls._parse_html(file_path)
        elif ext == ".json":
            content, meta = cls._parse_json(file_path)
        elif ext == ".csv":
            content, meta = cls._parse_csv(file_path)
        elif ext in (".txt",):
            content, meta = cls._parse_text(file_path)
        elif ext in (".py", ".js", ".ts", ".rs", ".go", ".java", ".c", ".cpp",
                     ".h", ".sh", ".bash", ".yml", ".yaml", ".toml", ".ini",
                     ".cfg", ".sql", ".graphql", ".xml"):
            content, meta = cls._parse_code(file_path)
        else:
            content, meta = cls._parse_text(file_path)

        title = cls._infer_title(file_path, content, meta)
        
        return {
            "title": title,
            "content": content,
            "format": ext.lstrip("."),
            "size": size,
            "parsed_at": datetime.now().isoformat(),
            "meta": meta,
        }

    @classmethod
    def _infer_title(cls, path: Path, content: str, meta: dict) -> str:
        """从内容或文件名推断标题."""
        # 1. Markdown 一级标题
        if path.suffix == ".md":
            for line in content.split("\n"):
                line = line.strip()
                if line.startswith("# ") and not line.startswith("## "):
                    return line[2:].strip()
        
        # 2. HTML title
        if "title" in meta:
            return meta["title"]
        
        # 3. PDF 标题
        if meta.get("pdf_title"):
            return meta["pdf_title"]
        
        # 4. DOCX 标题
        if meta.get("docx_title"):
            return meta["docx_title"]
        
        # 5. 文件名
        stem = path.stem
        # 将 kebab-case / snake_case 转为 Title Case
        title = re.sub(r'[-_]', ' ', stem).title()
        return title

    # ── 格式解析器 ──────────────────────────────────────────

    @classmethod
    def _parse_pdf(cls, path: Path) -> tuple[str, dict]:
        """解析 PDF — 保留段落结构."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            return f"*需要安装 PyPDF2: `pip install PyPDF2`*", {}

        reader = PdfReader(str(path))
        meta = {
            "pdf_pages": len(reader.pages),
            "pdf_title": (reader.metadata.title if reader.metadata and reader.metadata.title else ""),
            "pdf_author": (reader.metadata.author if reader.metadata and reader.metadata.author else ""),
        }

        sections = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                sections.append(text.strip())

        # 组装为美观的 Markdown
        lines = []
        if meta["pdf_title"]:
            lines.append(f"# {meta['pdf_title']}\n")
        if meta["pdf_author"]:
            lines.append(f"*{meta['pdf_author']}*\n")
        
        lines.append(f"> 📄 {meta['pdf_pages']} 页 | {cls._format_size(path)}\n")
        lines.append("---\n")
        
        for i, section in enumerate(sections):
            # 清理 PDF 提取的文本
            section = cls._clean_pdf_text(section)
            lines.append(section)
            if i < len(sections) - 1:
                lines.append("")  # 段落间空行

        return "\n".join(lines), meta

    @classmethod
    def _parse_docx(cls, path: Path) -> tuple[str, dict]:
        """解析 DOCX — 保留标题、段落、列表结构."""
        try:
            from docx import Document
        except ImportError:
            return f"*需要安装 python-docx: `pip install python-docx`*", {}

        doc = Document(str(path))
        meta = {
            "docx_title": "",
            "docx_author": "",
            "docx_paragraphs": len(doc.paragraphs),
        }
        
        # 提取元数据
        if doc.core_properties:
            cp = doc.core_properties
            meta["docx_title"] = cp.title or ""
            meta["docx_author"] = cp.author or ""

        lines = []
        if meta["docx_title"]:
            lines.append(f"# {meta['docx_title']}\n")
        if meta["docx_author"]:
            lines.append(f"*{meta['docx_author']}*\n")
        
        lines.append(f"> 📝 {meta['docx_paragraphs']} 段落 | {cls._format_size(path)}\n")
        lines.append("---\n")

        list_buffer = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if list_buffer:
                    lines.extend(list_buffer)
                    list_buffer = []
                lines.append("")
                continue

            style_name = (para.style.name or "").lower()
            
            # 标题
            if "heading" in style_name:
                if list_buffer:
                    lines.extend(list_buffer)
                    list_buffer = []
                level = 2
                for ch in style_name:
                    if ch.isdigit():
                        level = int(ch) + 1
                        break
                lines.append(f"{'#' * level} {text}\n")
            
            # 列表
            elif "list" in style_name:
                list_buffer.append(f"- {text}")
            
            # 引用
            elif "quote" in style_name or "intense quote" in style_name:
                if list_buffer:
                    lines.extend(list_buffer)
                    list_buffer = []
                lines.append(f"> {text}\n")
            
            # 普通段落
            else:
                if list_buffer:
                    lines.extend(list_buffer)
                    list_buffer = []
                
                # 处理内联格式
                formatted = cls._docx_inline_format(para)
                lines.append(formatted + "\n")

        if list_buffer:
            lines.extend(list_buffer)

        return "\n".join(lines), meta

    @classmethod
    def _docx_inline_format(cls, para) -> str:
        """提取 DOCX 段落的内联格式（粗体/斜体/代码）."""
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                parts.append(f"***{text}***")
            elif run.bold:
                parts.append(f"**{text}**")
            elif run.italic:
                parts.append(f"*{text}*")
            elif run.font.name and "mono" in run.font.name.lower():
                parts.append(f"`{text}`")
            else:
                parts.append(text)
        return "".join(parts) if parts else para.text

    @classmethod
    def _parse_markdown(cls, path: Path) -> tuple[str, dict]:
        """解析 Markdown — 直接读取."""
        content = path.read_text(encoding="utf-8", errors="replace")
        meta = {
            "md_headings": len(re.findall(r'^#{1,6}\s', content, re.MULTILINE)),
            "md_code_blocks": len(re.findall(r'```', content)) // 2,
            "md_links": len(re.findall(r'\[.*?\]\(.*?\)', content)),
        }
        return content, meta

    @classmethod
    def _parse_html(cls, path: Path) -> tuple[str, dict]:
        """解析 HTML — 转为 Markdown."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return path.read_text(encoding="utf-8", errors="replace"), {}

        html = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html, "html.parser")
        
        title = ""
        title_tag = soup.find("title")
        if title_tag:
            title = title_tag.get_text(strip=True)

        # 移除 script/style
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        body = soup.find("body") or soup
        meta = {"title": title}

        # 简单 HTML→Markdown 转换
        lines = []
        if title:
            lines.append(f"# {title}\n")
        
        for elem in body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "blockquote", "table"]):
            tag = elem.name
            text = elem.get_text(strip=True)
            
            if not text:
                continue
            
            if tag.startswith("h") and tag[1:].isdigit():
                level = int(tag[1:]) + 1
                lines.append(f"\n{'#' * level} {text}\n")
            elif tag == "p":
                lines.append(f"{text}\n")
            elif tag == "li":
                lines.append(f"- {text}")
            elif tag == "pre":
                lines.append(f"```\n{text}\n```\n")
            elif tag == "blockquote":
                for line in text.split("\n"):
                    lines.append(f"> {line}")
                lines.append("")

        return "\n".join(lines), meta

    @classmethod
    def _parse_json(cls, path: Path) -> tuple[str, dict]:
        """解析 JSON — 格式化显示."""
        content = path.read_text(encoding="utf-8", errors="replace")
        try:
            data = json.loads(content)
            formatted = json.dumps(data, indent=2, ensure_ascii=False)
            meta = {"json_keys": len(data) if isinstance(data, dict) else 0}
            
            lines = [
                f"> 📋 JSON | {cls._format_size(path)}\n",
                "---\n",
                f"```json\n{formatted}\n```",
            ]
            return "\n".join(lines), meta
        except json.JSONDecodeError:
            return f"```json\n{content}\n```", {}

    @classmethod
    def _parse_csv(cls, path: Path) -> tuple[str, dict]:
        """解析 CSV — 转为 Markdown 表格."""
        content = path.read_text(encoding="utf-8", errors="replace")
        rows = [line.split(",") for line in content.strip().split("\n") if line.strip()]
        
        if not rows:
            return "", {}

        # 构建 Markdown 表格
        lines = [
            f"> 📊 CSV 表格 | {len(rows)} 行 | {cls._format_size(path)}\n",
            "---\n",
        ]
        
        # 表头
        header = rows[0]
        lines.append("| " + " | ".join(h.strip() for h in header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        
        # 数据行（最多 50 行）
        for row in rows[1:51]:
            cells = [c.strip() for c in row]
            # 补齐列数
            while len(cells) < len(header):
                cells.append("")
            lines.append("| " + " | ".join(cells[:len(header)]) + " |")
        
        if len(rows) > 51:
            lines.append(f"\n*... 省略 {len(rows) - 51} 行*")

        meta = {"csv_rows": len(rows), "csv_columns": len(rows[0]) if rows else 0}
        return "\n".join(lines), meta

    @classmethod
    def _parse_text(cls, path: Path) -> tuple[str, dict]:
        """解析纯文本."""
        content = path.read_text(encoding="utf-8", errors="replace")
        meta = {"text_lines": content.count("\n") + 1}
        
        lines = [
            f"> 📝 文本文件 | {meta['text_lines']} 行 | {cls._format_size(path)}\n",
            "---\n",
            content,
        ]
        return "\n".join(lines), meta

    @classmethod
    def _parse_code(cls, path: Path) -> tuple[str, dict]:
        """解析代码文件 — 带语法高亮标记."""
        content = path.read_text(encoding="utf-8", errors="replace")
        lang = path.suffix.lstrip(".")
        
        # 语言映射
        lang_map = {
            "py": "python", "js": "javascript", "ts": "typescript",
            "rs": "rust", "go": "go", "java": "java",
            "c": "c", "cpp": "cpp", "h": "c",
            "sh": "bash", "bash": "bash",
            "yml": "yaml", "yaml": "yaml",
            "toml": "toml", "ini": "ini", "cfg": "ini",
            "sql": "sql", "graphql": "graphql", "xml": "xml",
        }
        code_lang = lang_map.get(lang, lang)
        
        meta = {
            "code_language": code_lang,
            "code_lines": content.count("\n") + 1,
        }
        
        # 截断过长的文件
        max_lines = 200
        lines_list = content.split("\n")
        if len(lines_list) > max_lines:
            truncated_content = "\n".join(lines_list[:max_lines])
            truncation_note = f"\n\n*... 截断，共 {meta['code_lines']} 行，仅显示前 {max_lines} 行*"
        else:
            truncated_content = content
            truncation_note = ""

        lines = [
            f"> 💻 {code_lang.upper()} | {meta['code_lines']} 行 | {cls._format_size(path)}\n",
            "---\n",
            f"```{code_lang}\n{truncated_content}\n```{truncation_note}",
        ]
        return "\n".join(lines), meta

    # ── 工具方法 ──────────────────────────────────────────

    @classmethod
    def _clean_pdf_text(cls, text: str) -> str:
        """清理 PDF 提取的文本."""
        # 合并被换行符打断的单词
        text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
        # 合并过短的换行
        lines = text.split("\n")
        merged = []
        for line in lines:
            line = line.strip()
            if not line:
                merged.append("")
            elif len(line) < 40 and not line.endswith((".", "!", "?", "。", "！", "？", "：", ":")):
                if merged and merged[-1] and not merged[-1].endswith("\n"):
                    merged[-1] += " " + line
                else:
                    merged.append(line)
            else:
                merged.append(line)
        return "\n".join(merged)

    @classmethod
    def _format_size(cls, path: Path) -> str:
        """格式化文件大小."""
        size = path.stat().st_size
        if size < 1024:
            return f"{size} B"
        elif size < 1024 * 1024:
            return f"{size / 1024:.1f} KB"
        else:
            return f"{size / (1024 * 1024):.1f} MB"
